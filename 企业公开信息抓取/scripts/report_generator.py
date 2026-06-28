#!/usr/bin/env python3
"""
report_generator.py — 生成Word经营分析报告

用法：
  python report_generator.py \
    --data D:\Workbuddy\企业披露信息抓取工作区\华天科技\01-财务报表\data.json \
    --company 华天科技 --code 002185 \
    --output D:\Workbuddy\企业披露信息抓取工作区\华天科技\02-分析报告\经营分析报告_20260521.docx
"""

import argparse
import json
import os
from datetime import datetime
from typing import Dict, Any, Optional


def fmt_num(val, unit="万元") -> str:
    """格式化数值"""
    if val is None:
        return "/"
    try:
        return f"{float(val):,.2f}{unit}"
    except (TypeError, ValueError):
        return str(val)


def fmt_pct(val) -> str:
    """格式化为百分比"""
    if val is None:
        return "/"
    try:
        return f"{float(val) * 100:.2f}%"
    except (TypeError, ValueError):
        return str(val)


def safe_div(a, b, default=None):
    try:
        return a / b if b else default
    except Exception:
        return default


def get_latest_years(data_by_year: dict, n: int = 3) -> list:
    """获取最新n年数据"""
    return sorted(data_by_year.keys())[-n:]


class ReportGenerator:
    def __init__(self, data_by_year: dict, company: str, code: str):
        self.data = data_by_year
        self.company = company
        self.code = code
        self.years = sorted(data_by_year.keys())
        self.latest_year = self.years[-1] if self.years else ""
        self.report_date = datetime.now().strftime("%Y年%m月%d日")

        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            self.Document = Document
            self.Pt = Pt
            self.Cm = Cm
            self.RGBColor = RGBColor
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("请安装 python-docx: pip install python-docx")

    def _latest_bs(self, scope="consolidated") -> dict:
        return self.data.get(self.latest_year, {}).get(f"balance_sheet_{scope}", {})

    def _latest_is(self, scope="consolidated") -> dict:
        return self.data.get(self.latest_year, {}).get(f"income_{scope}", {})

    def _latest_cf(self, scope="consolidated") -> dict:
        return self.data.get(self.latest_year, {}).get(f"cash_flow_{scope}", {})

    def _year_bs(self, year: str, scope="consolidated") -> dict:
        return self.data.get(year, {}).get(f"balance_sheet_{scope}", {})

    def _year_is(self, year: str, scope="consolidated") -> dict:
        return self.data.get(year, {}).get(f"income_{scope}", {})

    def _add_heading(self, doc, text: str, level: int):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = self.RGBColor(0x2B, 0x5B, 0xAE) if level == 1 else self.RGBColor(0, 0, 0)
        return h

    def _add_table(self, doc, data: list, headers: list, col_widths=None):
        """添加表格，data 为二维列表"""
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        table = doc.add_table(rows=1 + len(data), cols=len(headers))
        table.style = "Table Grid"

        # 表头
        header_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True

        # 数据行
        for ri, row in enumerate(data):
            tr = table.rows[ri + 1]
            for ci, val in enumerate(row):
                tr.cells[ci].text = str(val) if val is not None else "-"

        return table

    def generate(self, output_path: str):
        doc = self.Document()

        # ── 标题页 ──
        title = doc.add_heading(f"{self.company}\n经营情况分析报告", 0)
        title.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"报告日期：{self.report_date}").alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"数据来源：巨潮资讯网公开披露文件（{self.latest_year}年度报告）").alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # ── 第一章：企业基本情况 ──
        self._add_heading(doc, "一、企业基本情况", 1)
        self._write_basic_info(doc)

        # ── 第二章：财务状况分析 ──
        self._add_heading(doc, "二、财务状况分析", 1)
        self._write_financial_analysis(doc)

        # ── 第三章：经营分析 ──
        self._add_heading(doc, "三、经营分析", 1)
        self._write_operations_analysis(doc)

        # ── 第四章：重大事项 ──
        self._add_heading(doc, "四、重大事项", 1)
        self._write_major_events(doc)

        # ── 第五章：行业分析 ──
        self._add_heading(doc, "五、行业分析", 1)
        self._write_industry_analysis(doc)

        # ── 第六章：综合评价 ──
        self._add_heading(doc, "六、综合评价", 1)
        self._write_summary(doc)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"✅ 报告已保存: {output_path}")

    def _write_basic_info(self, doc):
        self._add_heading(doc, "（一）基本工商信息", 2)

        bs = self._latest_bs()
        equity = bs.get("total_equity", 0) or 0
        share_cap = bs.get("share_capital", bs.get("paid_in_capital", 0)) or 0

        data = [
            ["公司全称", self.company, "股票代码", self.code],
            ["注册资本（万元）", fmt_num(share_cap, ""), "成立时间", "（请补充）"],
            ["法定代表人", "（请补充）", "注册地", "（请补充）"],
            ["实际控制人", "（请补充）", "控股股东", "（请补充）"],
            ["主营业务", "（请补充）", "员工总数", "（请补充）"],
        ]
        table = doc.add_table(rows=len(data), cols=4)
        table.style = "Table Grid"
        for ri, row in enumerate(data):
            for ci, val in enumerate(row):
                c = table.rows[ri].cells[ci]
                c.text = str(val)
                if ci % 2 == 0:
                    c.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        self._add_heading(doc, "（二）主营业务简介", 2)
        doc.add_paragraph(
            f"{self.company}（股票代码：{self.code}）主营业务为（请根据年报补充）。"
            f"截至{self.latest_year}年末，公司资产总计 {fmt_num(bs.get('total_assets'))}，"
            f"所有者权益合计 {fmt_num(equity)}。"
        )

    def _write_financial_analysis(self, doc):
        """财务分析章节：资产负债 + 盈利 + 现金流"""
        self._add_heading(doc, "（一）资产负债情况", 2)

        # 多年对比表
        headers = ["科目（万元）"] + [f"{yr}年" for yr in self.years]
        rows = []
        key_items = [
            ("资产总计", "total_assets", "balance_sheet_consolidated"),
            ("负债合计", "total_liabilities", "balance_sheet_consolidated"),
            ("所有者权益合计", "total_equity", "balance_sheet_consolidated"),
            ("资产负债率", "__alr__", "balance_sheet_consolidated"),
            ("货币资金", "monetary_funds", "balance_sheet_consolidated"),
            ("短期借款", "short_term_loan", "balance_sheet_consolidated"),
            ("长期借款", "long_term_loan", "balance_sheet_consolidated"),
        ]

        for cn_name, key, section in key_items:
            row = [cn_name]
            for yr in self.years:
                yd = self.data.get(yr, {})
                if key == "__alr__":
                    ta = yd.get("balance_sheet_consolidated", {}).get("total_assets", 0) or 0
                    tl = yd.get("balance_sheet_consolidated", {}).get("total_liabilities", 0) or 0
                    val = fmt_pct(safe_div(tl, ta))
                else:
                    val = fmt_num(yd.get(section, {}).get(key), "")
                row.append(val)
            rows.append(row)

        self._add_table(doc, rows, headers)
        doc.add_paragraph()

        # 分析文字
        latest_bs = self._latest_bs()
        ta = latest_bs.get("total_assets", 0) or 0
        tl = latest_bs.get("total_liabilities", 0) or 0
        te = latest_bs.get("total_equity", 0) or 0
        alr = safe_div(tl, ta)

        doc.add_paragraph(
            f"{self.latest_year}年末，{self.company}资产总计{fmt_num(ta)}，较上年"
            f"（请补充增减幅度）；负债合计{fmt_num(tl)}，所有者权益合计{fmt_num(te)}，"
            f"资产负债率{fmt_pct(alr)}。"
        )

        # ── 盈利能力 ──
        self._add_heading(doc, "（二）盈利能力", 2)

        p_headers = ["科目（万元）"] + [f"{yr}年" for yr in self.years]
        p_rows = []
        p_items = [
            ("营业收入", "revenue", "income_consolidated"),
            ("营业成本", "cost_of_revenue", "income_consolidated"),
            ("毛利润", "__gross__", "income_consolidated"),
            ("净利润", "net_profit", "income_consolidated"),
            ("毛利率", "__gross_margin__", "income_consolidated"),
            ("净利率", "__net_margin__", "income_consolidated"),
        ]

        for cn_name, key, section in p_items:
            row = [cn_name]
            for yr in self.years:
                yd = self.data.get(yr, {})
                is_ = yd.get("income_consolidated", {})
                if key == "__gross__":
                    rev = is_.get("revenue", 0) or 0
                    cost = is_.get("cost_of_revenue", 0) or 0
                    val = fmt_num(rev - cost, "")
                elif key == "__gross_margin__":
                    rev = is_.get("revenue", 0) or 0
                    cost = is_.get("cost_of_revenue", 0) or 0
                    val = fmt_pct(safe_div(rev - cost, rev))
                elif key == "__net_margin__":
                    rev = is_.get("revenue", 0) or 0
                    np_ = is_.get("net_profit", 0) or 0
                    val = fmt_pct(safe_div(np_, rev))
                else:
                    val = fmt_num(is_.get(key), "")
                row.append(val)
            p_rows.append(row)

        self._add_table(doc, p_rows, p_headers)
        doc.add_paragraph()

        latest_is = self._latest_is()
        rev = latest_is.get("revenue", 0) or 0
        np_ = latest_is.get("net_profit", 0) or 0
        doc.add_paragraph(
            f"{self.latest_year}年，{self.company}实现营业收入{fmt_num(rev)}，"
            f"净利润{fmt_num(np_)}，净利率{fmt_pct(safe_div(np_, rev))}。"
        )

        # ── 现金流 ──
        self._add_heading(doc, "（三）现金流情况", 2)

        cf_headers = ["科目（万元）"] + [f"{yr}年" for yr in self.years]
        cf_rows = []
        cf_items = [
            ("经营活动现金流净额", "net_cash_from_operations"),
            ("投资活动现金流净额", "net_cash_from_investing"),
            ("筹资活动现金流净额", "net_cash_from_financing"),
            ("现金净增加额", "net_increase_in_cash"),
        ]

        for cn_name, key in cf_items:
            row = [cn_name]
            for yr in self.years:
                val = self.data.get(yr, {}).get("cash_flow_consolidated", {}).get(key)
                row.append(fmt_num(val, ""))
            cf_rows.append(row)

        self._add_table(doc, cf_rows, cf_headers)
        doc.add_paragraph()

        latest_cf = self._latest_cf()
        op_cf = latest_cf.get("net_cash_from_operations", 0) or 0
        doc.add_paragraph(
            f"{self.latest_year}年，{self.company}经营活动现金流量净额{fmt_num(op_cf)}，"
            f"经营现金流/净利润比率{fmt_pct(safe_div(op_cf, np_))}，现金流质量"
            f"{'良好' if op_cf > 0 and op_cf > np_ else '一般（请补充分析）'}。"
        )

    def _write_operations_analysis(self, doc):
        self._add_heading(doc, "（一）主营业务结构", 2)
        doc.add_paragraph(
            f"（请根据年报\"主营业务收入\"章节补充：分产品/分业务线收入、各业务毛利率及同比变化情况。）"
        )

        self._add_heading(doc, "（二）产能与产销", 2)
        doc.add_paragraph(
            f"（请根据年报补充：主要产品产能、产量、销量、库存情况，及在建工程/扩产计划。）"
        )

        self._add_heading(doc, "（三）核心竞争力", 2)
        doc.add_paragraph(
            f"（请根据年报\"公司核心竞争力分析\"章节补充：技术壁垒、市场份额、客户结构、品牌资质等。）"
        )

    def _write_major_events(self, doc):
        doc.add_paragraph(
            f"根据公开披露信息，{self.company}近期重大事项如下（如无则填写"经查询，报告期内无重大不利事项"）："
        )
        items = [
            "1. 重大资产重组/收购：（请补充）",
            "2. 股权融资（增发/配股）：（请补充）",
            "3. 关联交易：（请补充）",
            "4. 对外担保：（请补充）",
            "5. 重大诉讼/仲裁：（请补充）",
            "6. 前期会计差错更正：（请补充）",
            "7. 审计意见类型：（标准无保留意见 / 请补充）",
        ]
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    def _write_industry_analysis(self, doc):
        self._add_heading(doc, "（一）行业概况", 2)
        doc.add_paragraph(
            f"（请从最新年报\"报告期内公司所处行业情况\"章节提取：行业定义、监管分类、产业链结构。）"
        )

        self._add_heading(doc, "（二）竞争格局", 2)
        doc.add_paragraph(
            f"（请补充：主要竞争对手、市场集中度、公司市场地位，对比近2-3年竞争格局变化。）"
        )

        self._add_heading(doc, "（三）政策环境与发展趋势", 2)
        doc.add_paragraph(
            f"（请补充：行业相关政策、发展趋势、对公司经营的影响。）"
        )

    def _write_summary(self, doc):
        self._add_heading(doc, "（一）主要优势", 2)
        doc.add_paragraph("（请总结：公司核心优势，如市场地位、技术壁垒、客户资源、政策支持等。）")

        self._add_heading(doc, "（二）主要风险", 2)
        items = [
            "(1) 经营风险：（请补充）",
            "(2) 财务风险：（请补充，重点关注资产负债率、有息负债规模、流动性）",
            "(3) 行业风险：（请补充）",
            "(4) 政策风险：（请补充）",
        ]
        for item in items:
            doc.add_paragraph(item)

        # 发债企业额外节
        doc.add_paragraph()
        self._add_heading(doc, "（三）综合结论", 2)
        bs = self._latest_bs()
        is_ = self._latest_is()
        cf = self._latest_cf()
        ta = bs.get("total_assets", 0) or 0
        tl = bs.get("total_liabilities", 0) or 0
        alr = safe_div(tl, ta)
        rev = is_.get("revenue", 0) or 0
        np_ = is_.get("net_profit", 0) or 0
        op_cf = cf.get("net_cash_from_operations", 0) or 0

        doc.add_paragraph(
            f"{self.latest_year}年末，{self.company}资产总计{fmt_num(ta)}，"
            f"资产负债率{fmt_pct(alr)}；全年营业收入{fmt_num(rev)}，净利润{fmt_num(np_)}，"
            f"经营活动现金流净额{fmt_num(op_cf)}。整体来看，（请补充综合评价）。"
        )


# ─────────────────────────────────────────────
# 主函数
# ─────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="生成Word经营分析报告")
    parser.add_argument("--data", required=True, help="data.json 路径")
    parser.add_argument("--company", required=True, help="公司全称")
    parser.add_argument("--code", default="", help="股票代码")
    parser.add_argument("--output", required=True, help="输出 docx 路径")
    args = parser.parse_args()

    print("=== 经营分析报告生成器 ===")

    with open(args.data, "r", encoding="utf-8") as f:
        data_by_year = json.load(f)

    gen = ReportGenerator(data_by_year, args.company, args.code)
    gen.generate(args.output)


if __name__ == "__main__":
    main()
